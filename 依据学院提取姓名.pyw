from tkinter import *
import docx.enum.table
from docx import Document
import pandas as pd
from docx.shared import Cm
from docx.enum.text import *
from docx.enum.table import *
from docx.oxml.ns import qn

window = Tk()
window.title('拆分数据')
window.geometry('400x330')
var1 = StringVar()
var2 = StringVar()
var3 = StringVar()
var4 = StringVar()

L1 = Label(window, text='请输入需要提取文件所在完整目录(包含文件名和后缀)：', bg='orange', font=('Arial', 10), width=40, height=2)
L1.pack()
e1 = Entry(window, show=None, textvariable=var1, width=39, bd=4)
e1.pack()
L2 = Label(text='请输入文件存放位置（包含文件名和后缀）：', bg='orange', font=('Arial', 10), width=34, height=2)
L2.pack()
e2 = Entry(window, show=None, textvariable=var2, width=39, bd=4)
e2.pack()
L3 = Label(text='请输入文件中院系的表头：', bg='orange', font=('Arial', 12), width=30, height=2)
L3.pack()
e3 = Entry(window, show=None, textvariable=var3, width=39, bd=4)
e3.pack()
L4 = Label(text='请输入文件中姓名的表头：', bg='orange', font=('Arial', 12), width=30, height=2)
L4.pack()
e4 = Entry(window, show=None, textvariable=var4, width=39, bd=4)
e4.pack()


def data_division():
    file_path = e1.get()
    savePath = e2.get()
    colHeader = e3.get()
    nameHeader = e4.get()
    read_file = pd.read_excel(file_path)
    data_file = pd.DataFrame(read_file)
    coll_stu_list = []
    coo_stu_num = 0
    name_list = data_file[nameHeader].tolist()
    college_list = data_file[colHeader].tolist()
    college = college_list[0]
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    coll_num_list = []
    if len(name_list) != len(college_list):
        print('获取到的学生数和学院数对不上')
    for i in range(len(college_list)):
        if college != college_list[i]:
            if coo_stu_num != 0:
                p_head = document.add_heading('{}({})'.format(college, coo_stu_num), level=2)
                p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_head.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                temp_coo_stu_num = coo_stu_num
                row = 0
                if coo_stu_num % 7 == 0:
                    row = coo_stu_num // 7
                else:
                    row = int(coo_stu_num // 7) + 1
                tab1 = document.add_table(rows=row, cols=7)
                tab1.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                input_count = 0
                for tabel_index in range(row):
                    for table_colum in range(7):
                        cell = tab1.cell(tabel_index, table_colum)
                        cell.height = Cm(0.2)
                        if input_count < len(coll_stu_list):
                            cell.text = coll_stu_list[input_count]
                            # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                            input_count += 1
                if input_count != coo_stu_num:
                    print('学生列表输入有错误,输入数为{},学生列表数为{}'.format(input_count, coo_stu_num))
                # print(college_list[i])
                coll_num_list.append(coo_stu_num)
                coll_stu_list.clear()
                coo_stu_num = 0
                coll_stu_list.append(name_list[i])
                coo_stu_num += 1
                if i + 1 < len(college_list) - 1:
                    college = college_list[i]
            else:
                coll_stu_list.append(name_list[i])
                coo_stu_num += 1
        else:
            coll_stu_list.append(name_list[i])
            coo_stu_num += 1
            if i == len(college_list) - 1:
                p_head = document.add_heading('{}({})'.format(college, coo_stu_num), level=2)
                p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                temp_coo_stu_num = coo_stu_num
                row = 0
                if coo_stu_num % 7 == 0:
                    row = coo_stu_num // 7
                else:
                    row = int(coo_stu_num // 7) + 1
                table = document.add_table(rows=row, cols=7)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                input_count = 0
                for tabel_index in range(row):
                    for table_colum in range(7):
                        cell = table.cell(tabel_index, table_colum)
                        if input_count < len(coll_stu_list):
                            cell.text = coll_stu_list[input_count]
                            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                            input_count += 1
                if input_count != coo_stu_num:
                    print('学生列表输入有错误,输入数为{},学生列表数为{}'.format(input_count, coo_stu_num))
                # print(college_list[i])
                coll_num_list.append(coo_stu_num)
                coll_stu_list.clear()
                coo_stu_num = 0
                coll_stu_list.append(name_list[i])
                coo_stu_num += 1
                if i + 1 < len(college_list) - 1:
                    college = college_list[i + 1]
    document.save(savePath)
    print(coll_num_list)
    print('文档处理完毕！')


b = Button(window, text='开始', width=15, height=2, command=data_division)
b.pack()
window.mainloop()
